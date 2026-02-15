"""
End-to-end integration test script.

Creates a sample DOCX resume, uploads it via the API,
then runs a match against a sample job description.
Exercises the full pipeline: parse → NER → normalize → embed → graph → score → explain.

Usage:
    python scripts/e2e_test.py
"""
import io
import json
import sys

import httpx
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api/v1"


def create_sample_resume() -> bytes:
    """Generate a sample DOCX resume in memory."""
    doc = Document()
    doc.add_heading("John Doe — Senior ML Engineer", level=1)
    doc.add_paragraph(
        "Senior Machine Learning Engineer with 6+ years of experience "
        "building production ML systems. Proficient in Python, PyTorch, "
        "and TensorFlow. Experienced with NLP, computer vision, and deep learning."
    )
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(
        "Python, PyTorch, TensorFlow, scikit-learn, Pandas, NumPy, "
        "Docker, Kubernetes, AWS, PostgreSQL, Redis, FastAPI, "
        "Machine Learning, Deep Learning, NLP, Computer Vision, "
        "MLOps, CI/CD, Git, Linux, Apache Spark, Airflow"
    )
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(
        "ML Engineer at TechCorp (2020–2026)\n"
        "• Designed and deployed transformer-based NLP pipelines serving 10M+ requests/day\n"
        "• Built real-time recommendation engine using PyTorch and FAISS\n"
        "• Reduced model inference latency by 40% through TensorRT optimization\n"
        "• Led migration from monolith to microservices on Kubernetes"
    )
    doc.add_paragraph(
        "Data Scientist at DataLabs (2018–2020)\n"
        "• Built customer churn prediction models using scikit-learn and XGBoost\n"
        "• Developed data pipelines with Apache Spark and Airflow\n"
        "• Created dashboards with Tableau and Power BI"
    )
    doc.add_heading("Education", level=2)
    doc.add_paragraph("M.S. Computer Science — Stanford University, 2018")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    print("=" * 60)
    print("  Resume Job Fit Analyzer — End-to-End Test")
    print("=" * 60)

    # Step 1: Health check
    print("\n[1/3] Health check...")
    try:
        r = httpx.get(f"{BASE_URL}/health", timeout=10)
        health = r.json()
        print(f"  Status: {health['status']} | Version: {health['version']}")
    except httpx.ConnectError:
        print("  ERROR: Server not running. Start with: uvicorn app.main:app --port 8000")
        sys.exit(1)

    # Step 2: Upload resume
    print("\n[2/3] Uploading sample resume...")
    resume_bytes = create_sample_resume()
    r = httpx.post(
        f"{BASE_URL}/upload_resume",
        files={"file": ("john_doe_resume.docx", resume_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        timeout=120,  # Model loading can take time on first call
    )
    if r.status_code != 200:
        print(f"  ERROR: Upload failed ({r.status_code}): {r.text}")
        sys.exit(1)

    upload_data = r.json()
    resume_id = upload_data["resume_id"]
    print(f"  Resume ID: {resume_id}")
    print(f"  Skills found: {upload_data['skill_count']}")
    print(f"  Experience: {upload_data['experience_years']} years")

    # Step 3: Match against job description
    print("\n[3/3] Matching against Senior ML Engineer job...")
    match_payload = {
        "resume_id": resume_id,
        "job_description": {
            "title": "Senior ML Engineer",
            "description": (
                "We are looking for a Senior ML Engineer to join our AI team. "
                "You will design, build, and deploy production machine learning systems "
                "at scale. Must have strong experience with Python, deep learning frameworks, "
                "and cloud infrastructure."
            ),
            "required_skills": [
                "Python", "PyTorch", "Machine Learning", "Deep Learning",
                "Docker", "AWS", "NLP"
            ],
            "preferred_skills": [
                "Kubernetes", "TensorFlow", "MLOps", "FastAPI",
                "Apache Spark", "PostgreSQL"
            ],
            "min_experience_years": 4,
        },
    }

    r = httpx.post(
        f"{BASE_URL}/match",
        json=match_payload,
        timeout=120,
    )
    if r.status_code != 200:
        print(f"  ERROR: Match failed ({r.status_code}): {r.text}")
        sys.exit(1)

    result = r.json()

    print(f"\n{'=' * 60}")
    print(f"  MATCH RESULT")
    print(f"{'=' * 60}")
    print(f"  Overall Score: {result['overall_score']:.1%}")
    print(f"  Fit Label:     {result['fit_label']}")
    print(f"\n  Score Breakdown:")
    bd = result["score_breakdown"]
    print(f"    Semantic:    {bd['semantic_score']:.1%} (weight: {bd['semantic_weight']:.0%})")
    print(f"    Graph:       {bd['graph_score']:.1%} (weight: {bd['graph_weight']:.0%})")
    print(f"    Experience:  {bd['experience_score']:.1%} (weight: {bd['experience_weight']:.0%})")

    print(f"\n  Matched Skills ({len(result['matched_skills'])}):")
    for skill in sorted(result["matched_skills"], key=lambda x: x["similarity_score"], reverse=True):
        filled = int(skill["similarity_score"] * 10)
        bar = "#" * filled + "-" * (10 - filled)
        print(f"    [{bar}] {skill['similarity_score']:.1%}  {skill['skill']}")

    if result["missing_skills"]:
        print(f"\n  Missing Skills ({len(result['missing_skills'])}):")
        for skill in result["missing_skills"]:
            print(f"    x {skill}")

    print(f"\n{'=' * 60}")
    print(f"  Full Explanation:")
    print(f"{'=' * 60}")
    print(result["explanation"])

    print(f"\n{'=' * 60}")
    print(f"  END-TO-END TEST PASSED")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
